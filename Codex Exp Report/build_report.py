from __future__ import annotations

import math
import shutil
import subprocess
import textwrap
from dataclasses import dataclass
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt


SCRIPT_ROOT = Path(__file__).resolve().parent
THESIS_ROOT = SCRIPT_ROOT.parent
OUTPUT_ROOT = THESIS_ROOT / "Codex Exp Report"

ANTIGRAVITY_FIGURES = THESIS_ROOT / "final_experiments_report_antigravity" / "figures"
REPORTING_PACK = THESIS_ROOT / "artifacts" / "final_successful_runs" / "thesis_reporting_pack"

FIGURES_DIR = OUTPUT_ROOT / "figures"
DIAGRAMS_DIR = OUTPUT_ROOT / "diagrams"
TABLES_DIR = OUTPUT_ROOT / "tables"
DATA_DIR = OUTPUT_ROOT / "data"


@dataclass
class Summary:
    core_counts: pd.DataFrame
    core_runtime: pd.DataFrame
    core_winners: pd.DataFrame
    ablation_takeaways: pd.DataFrame
    point2_compliance: dict
    best_fertilization_group: dict
    best_fertilization_run: dict
    best_nonhier_group: dict
    best_nonhier_run: dict
    best_hier_group: dict
    best_overall_run: dict
    point1_random: dict
    point1_fixed: dict
    point1_random_stats: dict
    point2_best: pd.DataFrame
    point3_best_fixed: dict
    point3_best_random: dict
    point3_fixed_vs_08: dict
    point3_random_vs_12: dict
    study_inventory: dict


def ensure_dirs() -> None:
    for path in [OUTPUT_ROOT, FIGURES_DIR, DIAGRAMS_DIR, TABLES_DIR, DATA_DIR]:
        path.mkdir(parents=True, exist_ok=True)


def fmt_num(value: float, decimals: int = 2) -> str:
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return "n/a"
    return f"{value:,.{decimals}f}"


def fmt_int(value: float) -> str:
    if value is None or (isinstance(value, float) and math.isnan(value)):
        return "n/a"
    return f"{int(round(value)):,}"


def markdown_table(headers: list[str], rows: list[list[str]]) -> str:
    escaped_headers = [str(header).replace("|", "\\|") for header in headers]
    escaped_rows = [[str(cell).replace("|", "\\|") for cell in row] for row in rows]
    header_line = "| " + " | ".join(escaped_headers) + " |"
    divider_line = "| " + " | ".join(["---"] * len(headers)) + " |"
    body = "\n".join("| " + " | ".join(row) + " |" for row in escaped_rows)
    return "\n".join([header_line, divider_line, body])


def copy_reference_figures() -> None:
    figure_map = {
        ANTIGRAVITY_FIGURES / "final_113__leaderboard_primary_metric.png": FIGURES_DIR / "core_leaderboard.png",
        ANTIGRAVITY_FIGURES / "final_113__grouped_comparison.png": FIGURES_DIR / "core_best_runs.png",
        ANTIGRAVITY_FIGURES / "final_113__runtime_comparison.png": FIGURES_DIR / "core_runtime_comparison.png",
        ANTIGRAVITY_FIGURES / "final_42_ablation__point1_entropy_primary_metric.png": FIGURES_DIR / "ablation_entropy.png",
        ANTIGRAVITY_FIGURES / "final_42_ablation__point2_primary_comparison.png": FIGURES_DIR / "ablation_blocked_penalty.png",
        ANTIGRAVITY_FIGURES / "final_42_ablation__point2_thesis_compliance.png": FIGURES_DIR / "ablation_compliance.png",
        ANTIGRAVITY_FIGURES / "final_42_ablation__point3_cost_weight_primary_metric.png": FIGURES_DIR / "ablation_cost_weight.png",
        ANTIGRAVITY_FIGURES / "final_42_ablation__point3_cost_weight_paired_deltas.png": FIGURES_DIR / "ablation_cost_deltas.png",
        ANTIGRAVITY_FIGURES / "013_p2_a2c_fixed_weather_seed0_blockpen0__diagnostics_panel.png": FIGURES_DIR / "diagnostics_panel.png",
        ANTIGRAVITY_FIGURES / "013_p2_a2c_fixed_weather_seed0_blockpen0__weekly_npk_behavior.png": FIGURES_DIR / "weekly_npk_behavior.png",
        ANTIGRAVITY_FIGURES / "013_p2_a2c_fixed_weather_seed0_blockpen0__crop_decision_timeline.png": FIGURES_DIR / "crop_decision_timeline.png",
    }
    for src, dst in figure_map.items():
        shutil.copy2(src, dst)


def write_mermaid_sources() -> list[tuple[Path, Path]]:
    workflow = textwrap.dedent(
        """
        flowchart TD
            A["Pakistan weather, soil, and price inputs"] --> B["CyclesGym experiment environments"]
            B --> C{"Decision layer"}
            C --> C1["Weekly fertilization control"]
            C --> C2["Seasonal crop planning"]
            C --> C3["Hierarchical planning + fertilization"]
            C1 --> D["SB3 training, evaluation, and logging"]
            C2 --> D
            C3 --> D
            D --> E["Per-run summaries, diagnostics, and evaluation arrays"]
            E --> F["Grouped metrics, paired deltas, and runtime audits"]
            F --> G["Final experimentation chapter"]
        """
    ).strip()

    hierarchy = textwrap.dedent(
        """
        flowchart LR
            A["Start of agricultural year"] --> B["Select crop and planting window"]
            B --> C["Configure seasonal context"]
            C --> D["Take weekly NPK actions"]
            D --> E["Run CYCLES and parse outputs"]
            E --> F{"Season complete?"}
            F -- "No" --> D
            F -- "Yes" --> G["Aggregate return, cost, and compliance traces"]
            G --> H["Advance to next year"]
        """
    ).strip()

    mapping = [
        (DIAGRAMS_DIR / "experimental_workflow.mmd", DIAGRAMS_DIR / "experimental_workflow.png", workflow),
        (DIAGRAMS_DIR / "hierarchical_decision_flow.mmd", DIAGRAMS_DIR / "hierarchical_decision_flow.png", hierarchy),
    ]
    for src, _dst, content in mapping:
        src.write_text(content + "\n", encoding="utf-8")
    return [(src, dst) for src, dst, _content in mapping]


def render_mermaid(sources: list[tuple[Path, Path]]) -> None:
    mmdc = shutil.which("mmdc")
    if not mmdc:
        raise RuntimeError("Mermaid CLI `mmdc` is not available in PATH.")
    for src, dst in sources:
        if mmdc.lower().endswith(".ps1"):
            cmd = [
                "powershell",
                "-NoProfile",
                "-ExecutionPolicy",
                "Bypass",
                "-File",
                mmdc,
                "-i",
                str(src),
                "-o",
                str(dst),
                "-t",
                "neutral",
                "-b",
                "transparent",
            ]
        else:
            cmd = [mmdc, "-i", str(src), "-o", str(dst), "-t", "neutral", "-b", "transparent"]
        subprocess.run(cmd, check=True)


def save_data_inputs() -> None:
    input_files = {
        REPORTING_PACK / "final_113" / "tables" / "grouped" / "final_113__grouped_metrics.csv": DATA_DIR / "final_113__grouped_metrics.csv",
        REPORTING_PACK / "final_113" / "tables" / "grouped" / "final_113__run_catalog.csv": DATA_DIR / "final_113__run_catalog.csv",
        REPORTING_PACK / "final_113" / "tables" / "grouped" / "final_113__runtime_summary.csv": DATA_DIR / "final_113__runtime_summary.csv",
        REPORTING_PACK / "final_42_ablation" / "tables" / "grouped" / "final_42_ablation__point1_grouped_metrics.csv": DATA_DIR / "final_42_ablation__point1_grouped_metrics.csv",
        REPORTING_PACK / "final_42_ablation" / "tables" / "grouped" / "final_42_ablation__point1_paired_stats.csv": DATA_DIR / "final_42_ablation__point1_paired_stats.csv",
        REPORTING_PACK / "final_42_ablation" / "tables" / "grouped" / "final_42_ablation__point2_grouped_metrics.csv": DATA_DIR / "final_42_ablation__point2_grouped_metrics.csv",
        REPORTING_PACK / "final_42_ablation" / "tables" / "grouped" / "final_42_ablation__point3_grouped_metrics.csv": DATA_DIR / "final_42_ablation__point3_grouped_metrics.csv",
        REPORTING_PACK / "final_42_ablation" / "tables" / "grouped" / "final_42_ablation__point3_paired_stats.csv": DATA_DIR / "final_42_ablation__point3_paired_stats.csv",
    }
    for src, dst in input_files.items():
        shutil.copy2(src, dst)


def build_summary() -> Summary:
    core_grouped = pd.read_csv(DATA_DIR / "final_113__grouped_metrics.csv")
    core_runs = pd.read_csv(DATA_DIR / "final_113__run_catalog.csv")
    core_runtime = pd.read_csv(DATA_DIR / "final_113__runtime_summary.csv")
    point1 = pd.read_csv(DATA_DIR / "final_42_ablation__point1_grouped_metrics.csv")
    point1_stats = pd.read_csv(DATA_DIR / "final_42_ablation__point1_paired_stats.csv")
    point2 = pd.read_csv(DATA_DIR / "final_42_ablation__point2_grouped_metrics.csv")
    point3 = pd.read_csv(DATA_DIR / "final_42_ablation__point3_grouped_metrics.csv")
    point3_stats = pd.read_csv(DATA_DIR / "final_42_ablation__point3_paired_stats.csv")

    point2_compliance_values: list[float] = []
    point2_dir = REPORTING_PACK / "final_42_ablation" / "tables" / "per_run"
    for csv_path in point2_dir.glob("*season_window_compliance.csv"):
        df = pd.read_csv(csv_path)
        point2_compliance_values.extend(df["compliance_rate"].tolist())

    core_counts = (
        core_runs["report_group"]
        .value_counts()
        .rename_axis("report_group")
        .reset_index(name="run_count")
    )

    study_inventory = {
        "core_total_runs": int(len(core_runs)),
        "core_fertilization_runs": int((core_runs["domain"] == "fertilization").sum()),
        "core_crop_runs": int((core_runs["domain"] == "crop_planning").sum()),
        "ablation_total_runs": 42,
        "ablation_entropy_runs": 12,
        "ablation_shaping_runs": 12,
        "ablation_cost_runs": 18,
    }

    best_fertilization_group = (
        core_grouped[core_grouped["domain"] == "fertilization"]
        .sort_values("primary_metric_value_mean", ascending=False)
        .iloc[0]
        .to_dict()
    )
    best_nonhier_group = (
        core_grouped[core_grouped["report_group"] == "crop_planning_nonhier"]
        .sort_values("primary_metric_value_mean", ascending=False)
        .iloc[0]
        .to_dict()
    )
    best_hier_group = (
        core_grouped[core_grouped["report_group"] == "crop_planning_hierarchical_guarded_rerun"]
        .sort_values("primary_metric_value_mean", ascending=False)
        .iloc[0]
        .to_dict()
    )

    learned_core_runs = core_runs[core_runs["learned_run"] == True].copy()
    best_fertilization_run = (
        learned_core_runs[learned_core_runs["domain"] == "fertilization"]
        .sort_values("primary_metric_value", ascending=False)
        .iloc[0]
        .to_dict()
    )
    best_nonhier_run = (
        learned_core_runs[learned_core_runs["report_group"] == "crop_planning_nonhier"]
        .sort_values("primary_metric_value", ascending=False)
        .iloc[0]
        .to_dict()
    )
    best_overall_run = learned_core_runs.sort_values("primary_metric_value", ascending=False).iloc[0].to_dict()

    point1_fixed = point1[point1["weather_label"] == "fixed_weather"].sort_values("primary_metric_value__mean", ascending=False).iloc[0].to_dict()
    point1_random = point1[point1["weather_label"] == "random_weather"].sort_values("primary_metric_value__mean", ascending=False).iloc[0].to_dict()
    point1_random_stats = point1_stats[
        (point1_stats["metric"] == "deterministic_return")
        & (point1_stats["weather_label"] == "random_weather")
        & (point1_stats["ent_coef"] == 0.01)
    ].iloc[0].to_dict()

    point2_best = (
        point2.sort_values("deterministic_return", ascending=False)
        .groupby(["method", "weather_label"], as_index=False)
        .first()
        .sort_values(["method", "weather_label"])
    )

    point3_best_fixed = point3[point3["weather_label"] == "fixed_weather"].sort_values("primary_metric_value__mean", ascending=False).iloc[0].to_dict()
    point3_best_random = point3[point3["weather_label"] == "random_weather"].sort_values("primary_metric_value__mean", ascending=False).iloc[0].to_dict()
    point3_fixed_vs_08 = point3_stats[
        (point3_stats["metric"] == "deterministic_return")
        & (point3_stats["weather_label"] == "fixed_weather")
        & (point3_stats["nutrient_cost_weight"] == 0.8)
    ].iloc[0].to_dict()
    point3_random_vs_12 = point3_stats[
        (point3_stats["metric"] == "deterministic_return")
        & (point3_stats["weather_label"] == "random_weather")
        & (point3_stats["nutrient_cost_weight"] == 1.2)
    ].iloc[0].to_dict()

    core_winners = pd.DataFrame(
        [
            {
                "study_slice": "Fertilization grouped winner",
                "configuration": best_fertilization_group["group_key"],
                "metric": best_fertilization_group["primary_metric_name"],
                "value": fmt_num(best_fertilization_group["primary_metric_value_mean"]),
                "interval": f"[{fmt_num(best_fertilization_group['primary_metric_value_ci_low'])}, {fmt_num(best_fertilization_group['primary_metric_value_ci_high'])}]",
            },
            {
                "study_slice": "Fertilization best single run",
                "configuration": best_fertilization_run["run_slug"],
                "metric": best_fertilization_run["primary_metric_name"],
                "value": fmt_num(best_fertilization_run["primary_metric_value"]),
                "interval": "single run",
            },
            {
                "study_slice": "Non-hierarchical crop planning grouped winner",
                "configuration": best_nonhier_group["group_key"],
                "metric": best_nonhier_group["primary_metric_name"],
                "value": fmt_num(best_nonhier_group["primary_metric_value_mean"]),
                "interval": f"[{fmt_num(best_nonhier_group['primary_metric_value_ci_low'])}, {fmt_num(best_nonhier_group['primary_metric_value_ci_high'])}]",
            },
            {
                "study_slice": "Non-hierarchical crop planning best single run",
                "configuration": best_nonhier_run["run_slug"],
                "metric": best_nonhier_run["primary_metric_name"],
                "value": fmt_num(best_nonhier_run["primary_metric_value"], 3),
                "interval": "single run",
            },
            {
                "study_slice": "Hierarchical grouped winner",
                "configuration": best_hier_group["group_key"],
                "metric": best_hier_group["primary_metric_name"],
                "value": fmt_num(best_hier_group["primary_metric_value_mean"]),
                "interval": f"[{fmt_num(best_hier_group['primary_metric_value_ci_low'])}, {fmt_num(best_hier_group['primary_metric_value_ci_high'])}]",
            },
            {
                "study_slice": "Overall best single run",
                "configuration": best_overall_run["run_slug"],
                "metric": best_overall_run["primary_metric_name"],
                "value": fmt_num(best_overall_run["primary_metric_value"]),
                "interval": "single run",
            },
        ]
    )

    ablation_takeaways = pd.DataFrame(
        [
            {
                "ablation_point": "Entropy coefficient",
                "finding": "Entropy 0.00 was strongest in fixed weather, whereas entropy 0.01 was strongest in random weather.",
                "evidence": (
                    f"Fixed weather mean {fmt_num(point1_fixed['primary_metric_value__mean'])}; "
                    f"random weather mean {fmt_num(point1_random['primary_metric_value__mean'])}; "
                    f"paired random-weather delta {fmt_num(point1_random_stats['mean_delta'])} with p={point1_random_stats['p_value']:.4f}."
                ),
            },
            {
                "ablation_point": "Blocked-penalty shaping",
                "finding": "The shaping term benefited A2C but hurt PPO, while compliance stayed perfectly saturated.",
                "evidence": (
                    f"All {len(point2_compliance_values):,} recorded compliance-rate values were 1.0; "
                    f"A2C fixed-weather best penalty {point2_best[(point2_best['method'] == 'A2C') & (point2_best['weather_label'] == 'fixed_weather')].iloc[0]['blocked_penalty']:.2f}; "
                    f"PPO fixed-weather best penalty {point2_best[(point2_best['method'] == 'PPO') & (point2_best['weather_label'] == 'fixed_weather')].iloc[0]['blocked_penalty']:.2f}."
                ),
            },
            {
                "ablation_point": "Nutrient cost weight",
                "finding": "Cost weight 1.0 remained the best fixed-weather default; 1.2 was only marginally strongest in random weather.",
                "evidence": (
                    f"Fixed weather best mean {fmt_num(point3_best_fixed['primary_metric_value__mean'])} at weight {point3_best_fixed['nutrient_cost_weight']:.1f}; "
                    f"random weather best mean {fmt_num(point3_best_random['primary_metric_value__mean'])} at weight {point3_best_random['nutrient_cost_weight']:.1f}; "
                    f"random-weather 1.2 vs 1.0 delta {fmt_num(point3_random_vs_12['mean_delta'])} with p={point3_random_vs_12['p_value']:.4f}."
                ),
            },
        ]
    )

    return Summary(
        core_counts=core_counts,
        core_runtime=core_runtime,
        core_winners=core_winners,
        ablation_takeaways=ablation_takeaways,
        point2_compliance={
            "n_values": len(point2_compliance_values),
            "mean": float(sum(point2_compliance_values) / len(point2_compliance_values)),
            "min": float(min(point2_compliance_values)),
            "max": float(max(point2_compliance_values)),
        },
        best_fertilization_group=best_fertilization_group,
        best_fertilization_run=best_fertilization_run,
        best_nonhier_group=best_nonhier_group,
        best_nonhier_run=best_nonhier_run,
        best_hier_group=best_hier_group,
        best_overall_run=best_overall_run,
        point1_random=point1_random,
        point1_fixed=point1_fixed,
        point1_random_stats=point1_random_stats,
        point2_best=point2_best,
        point3_best_fixed=point3_best_fixed,
        point3_best_random=point3_best_random,
        point3_fixed_vs_08=point3_fixed_vs_08,
        point3_random_vs_12=point3_random_vs_12,
        study_inventory=study_inventory,
    )


def export_curated_tables(summary: Summary) -> None:
    summary.core_counts.to_csv(TABLES_DIR / "core_matrix_inventory.csv", index=False)
    summary.core_runtime.to_csv(TABLES_DIR / "core_runtime_summary.csv", index=False)
    summary.core_winners.to_csv(TABLES_DIR / "core_key_results.csv", index=False)
    summary.ablation_takeaways.to_csv(TABLES_DIR / "ablation_key_findings.csv", index=False)
    summary.point2_best.to_csv(TABLES_DIR / "point2_best_settings.csv", index=False)


def create_custom_figures(summary: Summary) -> None:
    plt.style.use("seaborn-v0_8-whitegrid")

    inventory = summary.core_counts.copy()
    inventory["label"] = inventory["report_group"].map(
        {
            "fertilization_core": "Fertilization core",
            "crop_planning_nonhier": "Crop planning\n(non-hierarchical)",
            "crop_planning_hierarchical_guarded_rerun": "Crop planning\n(hierarchical)",
            "fertilization_dqn_rerun": "Fertilization DQN",
            "crop_planning_dqn_rerun": "Crop planning DQN",
            "fertilization_baseline": "Baseline only",
        }
    )

    fig, ax = plt.subplots(figsize=(10, 5))
    bars = ax.bar(inventory["label"], inventory["run_count"], color="#1f4e79")
    ax.set_title("Core matrix inventory by study slice", fontsize=14, weight="bold")
    ax.set_ylabel("Number of runs")
    ax.set_xlabel("")
    ax.bar_label(bars, padding=3)
    plt.tight_layout()
    plt.savefig(FIGURES_DIR / "study_inventory.png", dpi=300)
    plt.close(fig)

    runtime = summary.core_runtime.copy()
    runtime["label"] = runtime["report_group"].map(
        {
            "fertilization_core": "Fertilization core",
            "crop_planning_nonhier": "Crop planning\n(non-hierarchical)",
            "crop_planning_hierarchical_guarded_rerun": "Crop planning\n(hierarchical)",
            "fertilization_dqn_rerun": "Fertilization DQN",
            "crop_planning_dqn_rerun": "Crop planning DQN",
        }
    )
    runtime = runtime.dropna(subset=["label"])
    runtime["hours"] = runtime["runtime_seconds_mean"] / 3600.0

    fig, ax = plt.subplots(figsize=(10, 5))
    bars = ax.bar(runtime["label"], runtime["hours"], color="#b85c38")
    ax.set_title("Mean runtime by report group", fontsize=14, weight="bold")
    ax.set_ylabel("Mean runtime (hours)")
    ax.set_xlabel("")
    ax.bar_label(bars, labels=[f"{val:.2f}" for val in runtime["hours"]], padding=3)
    plt.tight_layout()
    plt.savefig(FIGURES_DIR / "runtime_tradeoffs.png", dpi=300)
    plt.close(fig)


def add_doc_title(document: Document, title: str, subtitle: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(18)
    run.font.name = "Times New Roman"

    p2 = document.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = p2.add_run(subtitle)
    run2.italic = True
    run2.font.size = Pt(11)
    run2.font.name = "Times New Roman"


def set_document_style(document: Document) -> None:
    section = document.sections[0]
    section.page_width = Inches(8.27)
    section.page_height = Inches(11.69)
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    style = document.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(12)

    for heading_name, size in [("Heading 1", 15), ("Heading 2", 13), ("Heading 3", 12)]:
        heading = document.styles[heading_name]
        heading.font.name = "Times New Roman"
        heading.font.size = Pt(size)
        heading.font.bold = True


def add_body_paragraph(document: Document, text: str) -> None:
    paragraph = document.add_paragraph(text)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY


def add_bullets(document: Document, items: list[str]) -> None:
    for item in items:
        paragraph = document.add_paragraph(style="List Bullet")
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        run = paragraph.add_run(item)
        run.font.name = "Times New Roman"
        run.font.size = Pt(12)


def add_table(document: Document, dataframe: pd.DataFrame, caption: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(caption)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)

    table = document.add_table(rows=1, cols=len(dataframe.columns))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr_cells = table.rows[0].cells
    for idx, column in enumerate(dataframe.columns):
        hdr_cells[idx].text = str(column)
    for _idx, row in dataframe.iterrows():
        cells = table.add_row().cells
        for cell_index, value in enumerate(row):
            cells[cell_index].text = str(value)
    document.add_paragraph()


def add_figure(document: Document, image_path: Path, caption: str, width_inches: float = 6.2) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width_inches))

    cap = document.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cap.add_run(caption)
    run.italic = True
    run.font.name = "Times New Roman"
    run.font.size = Pt(11)


def build_docx(summary: Summary) -> None:
    document = Document()
    set_document_style(document)
    document.core_properties.title = "Final Experimentation Report"
    document.core_properties.author = "Codex"

    add_doc_title(
        document,
        "Final Experimentation Report",
        "Pakistan-adapted reinforcement learning experiments for long-horizon crop management",
    )

    document.add_heading("Abstract", level=1)
    add_body_paragraph(
        document,
        (
            "This report consolidates the finalized experimentation evidence for the thesis on cost-aware "
            "reinforcement learning for agricultural resource allocation in a Pakistan-adapted CyclesGym stack. "
            f"It integrates a {summary.study_inventory['core_total_runs']}-run core matrix and a "
            f"{summary.study_inventory['ablation_total_runs']}-run ablation suite into a single thesis-ready narrative. "
            f"The strongest grouped result in fertilization was observed for {summary.best_fertilization_group['group_key']}, "
            f"which reached a mean primary metric of {fmt_num(summary.best_fertilization_group['primary_metric_value_mean'])}. "
            f"The strongest grouped result in hierarchical crop planning was {summary.best_hier_group['group_key']}, "
            f"with a mean deterministic return of {fmt_num(summary.best_hier_group['primary_metric_value_mean'])}, while the "
            f"best overall single run was {summary.best_overall_run['run_slug']} at "
            f"{fmt_num(summary.best_overall_run['primary_metric_value'])}. "
            f"The ablation suite shows a clear interaction between entropy and environmental stochasticity: "
            f"zero entropy was strongest in fixed weather, whereas entropy 0.01 improved random-weather deterministic return by "
            f"{fmt_num(summary.point1_random_stats['mean_delta'])} with p={summary.point1_random_stats['p_value']:.4f}. "
            f"Blocked-penalty shaping should be interpreted descriptively rather than inferentially because it uses single-seed "
            f"comparisons, but every recorded compliance-rate value remained 1.0, implying reward reshaping rather than constraint enforcement. "
            f"Across the nutrient-cost ablation, weight 1.0 remained the safest default despite a marginal random-weather edge for 1.2."
        ),
    )

    document.add_heading("1. Experimental objectives and evidence base", level=1)
    add_body_paragraph(
        document,
        (
            "The purpose of this report is to convert the available experimentation artifacts into a final, academically coherent "
            "results chapter for the thesis. The report is grounded in the thesis framing from the reference LaTeX archive and in "
            "the audited grouped CSV packs that back the antigravity report. The empirical evidence consists of "
            f"{summary.study_inventory['core_fertilization_runs']} fertilization runs and {summary.study_inventory['core_crop_runs']} crop-planning runs "
            "in the core matrix, together with three focused ablation branches covering entropy, blocked-penalty shaping, and nutrient-cost weighting."
        ),
    )
    add_figure(
        document,
        FIGURES_DIR / "study_inventory.png",
        "Figure 1. Inventory of the completed core matrix by study slice.",
    )
    add_table(
        document,
        pd.DataFrame(
            {
                "Report group": summary.core_counts["report_group"].map(
                    {
                        "fertilization_core": "Fertilization core",
                        "crop_planning_nonhier": "Crop planning (non-hierarchical)",
                        "crop_planning_hierarchical_guarded_rerun": "Crop planning (hierarchical guarded reruns)",
                        "fertilization_dqn_rerun": "Fertilization DQN reruns",
                        "crop_planning_dqn_rerun": "Crop planning DQN reruns",
                        "fertilization_baseline": "Baseline only",
                    }
                ),
                "Runs": summary.core_counts["run_count"].map(fmt_int),
            }
        ),
        "Table 1. Distribution of the completed core matrix.",
    )

    document.add_heading("2. Experimental protocol", level=1)
    add_body_paragraph(
        document,
        (
            "The experimentation stack evaluates three decision layers: weekly fertilization control, seasonal crop planning, and a "
            "hierarchical formulation that combines yearly crop choice with within-season nutrient actions. The underlying thesis framing "
            "is Pakistan-oriented and cost-aware: weather, soil, and price inputs are localized, but the active experimental configuration "
            "remains simulation-based and tied to the currently stabilized maize-soy setup. PPO and A2C are the main algorithms; DQN is included "
            "only as a descriptive reference rather than a seed-balanced competitor."
        ),
    )
    add_figure(
        document,
        DIAGRAMS_DIR / "experimental_workflow.png",
        "Figure 2. End-to-end experiment workflow from localized inputs to thesis-ready grouped metrics.",
    )
    add_figure(
        document,
        DIAGRAMS_DIR / "hierarchical_decision_flow.png",
        "Figure 3. Two-timescale decision structure used by the hierarchical crop-planning formulation.",
    )
    add_body_paragraph(
        document,
        (
            "Metrics were interpreted in layers. Grouped primary metrics were used for ranking, deterministic and stochastic returns for "
            "behavioral interpretation, and runtime and compliance outputs for operational context. This layered reading is important because "
            "a thesis defense requires more than a leaderboard; it requires a clear explanation of how the policy behaves and under what constraints."
        ),
    )

    document.add_heading("3. Core matrix results", level=1)
    add_body_paragraph(
        document,
        (
            "The core matrix shows that the strongest absolute returns came from hierarchical crop-planning reruns, whereas the strongest grouped "
            "fertilization result remained an A2C configuration under deterministic weather and a long training horizon. This split matters because it "
            "shows that the best-performing method depends on both the decision layer and the degree of environmental stochasticity."
        ),
    )
    add_table(document, summary.core_winners, "Table 2. Key grouped winners and best single runs extracted from the core matrix.")
    add_figure(
        document,
        FIGURES_DIR / "core_leaderboard.png",
        "Figure 4. Grouped leaderboard from the Final 113 core matrix.",
    )
    add_figure(
        document,
        FIGURES_DIR / "core_best_runs.png",
        "Figure 5. Best single-run outcomes in the core matrix, dominated by hierarchical crop-planning reruns.",
    )

    document.add_heading("3.1 Fertilization", level=2)
    add_body_paragraph(
        document,
        (
            f"The strongest grouped fertilization result was {summary.best_fertilization_group['group_key']}, which achieved a mean deterministic "
            f"return of {fmt_num(summary.best_fertilization_group['primary_metric_value_mean'])} with a 95% interval of "
            f"[{fmt_num(summary.best_fertilization_group['primary_metric_value_ci_low'])}, "
            f"{fmt_num(summary.best_fertilization_group['primary_metric_value_ci_high'])}] across three seeds. "
            f"The strongest single fertilization run was {summary.best_fertilization_run['run_slug']} at "
            f"{fmt_num(summary.best_fertilization_run['primary_metric_value'])}. "
            "Substantively, this suggests that A2C remains competitive when the environment is deterministic and the budget is long, while PPO can still "
            "produce seed-level peaks under shorter fixed-weather configurations."
        ),
    )

    document.add_heading("3.2 Non-hierarchical crop planning", level=2)
    add_body_paragraph(
        document,
        (
            f"For non-hierarchical crop planning, the strongest grouped result was {summary.best_nonhier_group['group_key']} with a mean evaluation reward "
            f"of {fmt_num(summary.best_nonhier_group['primary_metric_value_mean'])}. "
            f"The strongest single non-hierarchical run was {summary.best_nonhier_run['run_slug']} at "
            f"{fmt_num(summary.best_nonhier_run['primary_metric_value'], 3)}. "
            "This ranking favors PPO for grouped stability, but it also shows that A2C can still discover favorable single-seed trajectories. The correct "
            "interpretation is therefore not that one method dominates everywhere, but that PPO offers the stronger default for repeatable non-hierarchical "
            "crop-planning performance."
        ),
    )

    document.add_heading("3.3 Hierarchical crop planning and runtime tradeoffs", level=2)
    add_body_paragraph(
        document,
        (
            f"The strongest grouped hierarchical result was {summary.best_hier_group['group_key']}, which reached "
            f"{fmt_num(summary.best_hier_group['primary_metric_value_mean'])} with a 95% interval of "
            f"[{fmt_num(summary.best_hier_group['primary_metric_value_ci_low'])}, {fmt_num(summary.best_hier_group['primary_metric_value_ci_high'])}]. "
            f"The best overall run in the entire evidence pack was {summary.best_overall_run['run_slug']} at "
            f"{fmt_num(summary.best_overall_run['primary_metric_value'])}. "
            "These results indicate that once yearly crop choice is coupled to within-season fertilization, the value of hierarchical structure is large enough "
            "to dominate the experiment leaderboard. The tradeoff is computational: hierarchical reruns are materially slower than non-hierarchical crop-planning "
            "jobs and fertilization runs."
        ),
    )
    add_figure(
        document,
        FIGURES_DIR / "runtime_tradeoffs.png",
        "Figure 6. Mean runtime by report group, showing the heavier cost of hierarchical reruns.",
    )
    add_figure(
        document,
        FIGURES_DIR / "core_runtime_comparison.png",
        "Figure 7. Distributional runtime comparison from the core matrix export pack.",
    )

    document.add_heading("4. Ablation results", level=1)
    add_body_paragraph(
        document,
        (
            "The 42-run ablation suite should be read as a parameter-isolation study rather than as a second leaderboard. It answers three narrower questions: "
            "how much exploration pressure is needed, whether blocked-penalty shaping adds value beyond existing guards, and what nutrient-cost weight behaves "
            "most safely across weather regimes."
        ),
    )
    add_table(document, summary.ablation_takeaways, "Table 3. High-level conclusions from the ablation suite.")

    document.add_heading("4.1 Entropy coefficient", level=2)
    add_body_paragraph(
        document,
        (
            f"In fixed weather, the best grouped entropy setting was ent_coef={summary.point1_fixed['ent_coef']:.2f} with a mean primary metric of "
            f"{fmt_num(summary.point1_fixed['primary_metric_value__mean'])}. In random weather, the best grouped setting was ent_coef="
            f"{summary.point1_random['ent_coef']:.2f} with a mean of {fmt_num(summary.point1_random['primary_metric_value__mean'])}. "
            f"The paired random-weather deterministic-return delta for ent_coef 0.01 relative to 0.00 was "
            f"{fmt_num(summary.point1_random_stats['mean_delta'])} with p={summary.point1_random_stats['p_value']:.4f}. "
            "This is the clearest inferential result in the ablation pack: entropy hurts convergence in deterministic regimes but becomes helpful once weather "
            "uncertainty is reintroduced."
        ),
    )
    add_figure(
        document,
        FIGURES_DIR / "ablation_entropy.png",
        "Figure 8. Entropy ablation showing the reversal between fixed and random weather.",
    )

    document.add_heading("4.2 Blocked-penalty shaping", level=2)
    add_body_paragraph(
        document,
        (
            "The blocked-penalty ablation is descriptive because each configuration is represented by a single seed. Its most defensible finding is not the exact "
            "size of the return change, but the mechanism behind it: every recorded season-window compliance value remained 1.0, so the penalty did not improve rule "
            "adherence because adherence was already saturated. Instead, it changed the reward landscape. In A2C, moderate or strong penalty values improved the "
            "deterministic return. In PPO, the zero-penalty configuration remained best under both weather regimes."
        ),
    )
    add_figure(
        document,
        FIGURES_DIR / "ablation_blocked_penalty.png",
        "Figure 9. Deterministic-return comparison across blocked-penalty settings.",
    )
    add_figure(
        document,
        FIGURES_DIR / "ablation_compliance.png",
        "Figure 10. Compliance remained saturated across blocked-penalty settings.",
    )

    document.add_heading("4.3 Nutrient-cost weight", level=2)
    add_body_paragraph(
        document,
        (
            f"The best fixed-weather cost weight was {summary.point3_best_fixed['nutrient_cost_weight']:.1f}, which achieved a mean primary metric of "
            f"{fmt_num(summary.point3_best_fixed['primary_metric_value__mean'])}. "
            f"In random weather, the highest grouped mean was observed at weight {summary.point3_best_random['nutrient_cost_weight']:.1f}, but the margin over "
            f"weight 1.0 was only {fmt_num(summary.point3_random_vs_12['mean_delta'])} and was not statistically persuasive "
            f"(p={summary.point3_random_vs_12['p_value']:.4f}). By contrast, under fixed weather the drop from 1.0 to 0.8 was "
            f"{fmt_num(abs(summary.point3_fixed_vs_08['mean_delta']))}. The academically careful conclusion is therefore that weight 1.0 remains the safest thesis "
            "default, while 1.2 may be explored in random weather if a slightly more conservative nutrient policy is desired."
        ),
    )
    add_figure(
        document,
        FIGURES_DIR / "ablation_cost_weight.png",
        "Figure 11. Grouped nutrient-cost-weight performance under fixed and random weather.",
    )
    add_figure(
        document,
        FIGURES_DIR / "ablation_cost_deltas.png",
        "Figure 12. Paired delta view for the nutrient-cost-weight ablation.",
    )

    document.add_heading("5. Behavioral diagnostics", level=1)
    add_body_paragraph(
        document,
        (
            "Grouped means explain which settings won, but the per-run diagnostics explain how those wins were realized. The extracted diagnostics show that the "
            "learning curves in successful runs become progressively smoother, that nutrient application is not uniformly distributed across the season, and that the "
            "hierarchical policies produce legible crop-decision traces. These properties matter because they make the experimentation chapter more interpretable than "
            "a pure return table."
        ),
    )
    add_figure(
        document,
        FIGURES_DIR / "diagnostics_panel.png",
        "Figure 13. Representative training diagnostics from the blocked-penalty study.",
    )
    add_figure(
        document,
        FIGURES_DIR / "weekly_npk_behavior.png",
        "Figure 14. Weekly NPK behavior from a representative hierarchical run.",
    )
    add_figure(
        document,
        FIGURES_DIR / "crop_decision_timeline.png",
        "Figure 15. Crop-decision timeline illustrating the policy's yearly planning structure.",
    )

    document.add_heading("6. Discussion and validity boundaries", level=1)
    add_body_paragraph(
        document,
        (
            "Three conclusions can be defended cleanly from this evidence pack. First, hierarchical crop planning produces the strongest absolute returns. Second, "
            "fertilization behavior is method-sensitive: A2C is highly competitive in deterministic fixed-weather settings, while PPO remains the more robust general "
            "choice once broader stochasticity and hierarchy are introduced. Third, the best hyperparameters are regime-dependent: entropy and cost weighting cannot be "
            "set intelligently without considering weather uncertainty."
        ),
    )
    add_bullets(
        document,
        [
            "Simulation is the evaluation medium throughout; none of the conclusions should be presented as field-validated agronomy.",
            "The active crop-planning and hierarchical experiments remain tied to the current stabilized maize-soy configuration, even though the broader thesis framing is Pakistan-oriented.",
            "Most grouped comparisons use only three seeds, and the blocked-penalty ablation uses one seed per configuration, so effect size should be interpreted more cautiously than direction of effect.",
            "DQN was only explored in small reruns and therefore remains a descriptive reference rather than a full competitor in the main ranking.",
        ],
    )
    add_body_paragraph(
        document,
        (
            "Within those limits, the practical thesis recommendation is clear: use hierarchical PPO as the strongest overall configuration, keep entropy at 0.00 for "
            "fixed weather and 0.01 for random weather, retain nutrient cost weight 1.0 as the default, and apply blocked-penalty shaping only with explicit awareness "
            "that its benefit appears algorithm-specific rather than universal."
        ),
    )

    document.add_heading("7. Conclusion", level=1)
    add_body_paragraph(
        document,
        (
            "The finalized experimentation evidence supports a coherent thesis story. The Pakistan-adapted RL stack is not merely run-ready; it already yields a "
            "structured body of results from which method, weather, and design choices can be interpreted with appropriate caution. Hierarchical crop planning is the "
            "strongest empirical theme, fixed-weather fertilization remains a competitive niche for A2C, and the ablation suite clarifies which configuration defaults "
            "should carry forward into the final thesis narrative. The main scientific value of this chapter lies not only in the scores themselves, but also in the "
            "fact that the report can now explain those scores in terms of architecture, stochasticity, reward shaping, and operational behavior."
        ),
    )

    document.save(OUTPUT_ROOT / "final_experimentation_report.docx")


def build_markdown(summary: Summary) -> None:
    inventory_rows = [
        [row["report_group"].replace("_", " "), fmt_int(row["run_count"])]
        for _idx, row in summary.core_counts.iterrows()
    ]

    winners_rows = [[str(item) for item in row] for row in summary.core_winners.values.tolist()]
    ablation_rows = [[str(item) for item in row] for row in summary.ablation_takeaways.values.tolist()]

    markdown = f"""# Final Experimentation Report

*Pakistan-adapted reinforcement learning experiments for long-horizon crop management*

## Abstract

This report consolidates the finalized experimentation evidence for the thesis on cost-aware reinforcement learning for agricultural resource allocation in a Pakistan-adapted CyclesGym stack. It integrates a {summary.study_inventory['core_total_runs']}-run core matrix and a {summary.study_inventory['ablation_total_runs']}-run ablation suite into a single thesis-ready narrative. The strongest grouped fertilization result was observed for `{summary.best_fertilization_group['group_key']}`, which reached a mean primary metric of {fmt_num(summary.best_fertilization_group['primary_metric_value_mean'])}. The strongest grouped hierarchical result was `{summary.best_hier_group['group_key']}` at {fmt_num(summary.best_hier_group['primary_metric_value_mean'])}, while the best overall single run was `{summary.best_overall_run['run_slug']}` at {fmt_num(summary.best_overall_run['primary_metric_value'])}. In the ablation suite, entropy 0.00 remained strongest in fixed weather, entropy 0.01 improved random-weather deterministic return by {fmt_num(summary.point1_random_stats['mean_delta'])} with `p={summary.point1_random_stats['p_value']:.4f}`, blocked-penalty shaping acted as reward reshaping because compliance was saturated at 1.0, and nutrient cost weight 1.0 remained the safest overall default.

## 1. Experimental objectives and evidence base

The purpose of this report is to convert the available experimentation artifacts into a final, academically coherent experimentation chapter. The report is grounded in the thesis framing recovered from the reference LaTeX archive and in the audited grouped CSV packs behind the antigravity report. The empirical evidence consists of {summary.study_inventory['core_fertilization_runs']} fertilization runs and {summary.study_inventory['core_crop_runs']} crop-planning runs in the core matrix, together with three focused ablation branches covering entropy, blocked-penalty shaping, and nutrient-cost weighting.

![Study inventory](figures/study_inventory.png)

*Figure 1. Inventory of the completed core matrix by study slice.*

{markdown_table(['Report group', 'Runs'], inventory_rows)}

## 2. Experimental protocol

The experimentation stack evaluates three decision layers: weekly fertilization control, seasonal crop planning, and a hierarchical formulation that combines yearly crop choice with within-season nutrient actions. The underlying thesis framing is Pakistan-oriented and cost-aware: weather, soil, and price inputs are localized, but the active experimental configuration remains simulation-based and tied to the currently stabilized maize-soy setup. PPO and A2C are the main algorithms; DQN is included only as a descriptive reference rather than a seed-balanced competitor.

![Experimental workflow](diagrams/experimental_workflow.png)

*Figure 2. End-to-end experiment workflow from localized inputs to thesis-ready grouped metrics.*

![Hierarchical decision flow](diagrams/hierarchical_decision_flow.png)

*Figure 3. Two-timescale decision structure used by the hierarchical crop-planning formulation.*

Metrics were interpreted in layers. Grouped primary metrics were used for ranking, deterministic and stochastic returns for behavioral interpretation, and runtime and compliance outputs for operational context. This layered reading is important because a thesis defense requires more than a leaderboard; it requires a clear explanation of how the policy behaves and under what constraints.

## 3. Core matrix results

The core matrix shows that the strongest absolute returns came from hierarchical crop-planning reruns, whereas the strongest grouped fertilization result remained an A2C configuration under deterministic weather and a long training horizon. This split matters because it shows that the best-performing method depends on both the decision layer and the degree of environmental stochasticity.

{markdown_table(list(summary.core_winners.columns), winners_rows)}

![Core leaderboard](figures/core_leaderboard.png)

*Figure 4. Grouped leaderboard from the Final 113 core matrix.*

![Best runs](figures/core_best_runs.png)

*Figure 5. Best single-run outcomes in the core matrix, dominated by hierarchical crop-planning reruns.*

### 3.1 Fertilization

The strongest grouped fertilization result was `{summary.best_fertilization_group['group_key']}`, which achieved a mean deterministic return of {fmt_num(summary.best_fertilization_group['primary_metric_value_mean'])} with a 95% interval of [{fmt_num(summary.best_fertilization_group['primary_metric_value_ci_low'])}, {fmt_num(summary.best_fertilization_group['primary_metric_value_ci_high'])}] across three seeds. The strongest single fertilization run was `{summary.best_fertilization_run['run_slug']}` at {fmt_num(summary.best_fertilization_run['primary_metric_value'])}. Substantively, this suggests that A2C remains competitive when the environment is deterministic and the budget is long, while PPO can still produce seed-level peaks under shorter fixed-weather configurations.

### 3.2 Non-hierarchical crop planning

For non-hierarchical crop planning, the strongest grouped result was `{summary.best_nonhier_group['group_key']}` with a mean evaluation reward of {fmt_num(summary.best_nonhier_group['primary_metric_value_mean'])}. The strongest single non-hierarchical run was `{summary.best_nonhier_run['run_slug']}` at {fmt_num(summary.best_nonhier_run['primary_metric_value'], 3)}. This ranking favors PPO for grouped stability, but it also shows that A2C can still discover favorable single-seed trajectories. The correct interpretation is therefore not that one method dominates everywhere, but that PPO offers the stronger default for repeatable non-hierarchical crop-planning performance.

### 3.3 Hierarchical crop planning and runtime tradeoffs

The strongest grouped hierarchical result was `{summary.best_hier_group['group_key']}`, which reached {fmt_num(summary.best_hier_group['primary_metric_value_mean'])} with a 95% interval of [{fmt_num(summary.best_hier_group['primary_metric_value_ci_low'])}, {fmt_num(summary.best_hier_group['primary_metric_value_ci_high'])}]. The best overall run in the entire evidence pack was `{summary.best_overall_run['run_slug']}` at {fmt_num(summary.best_overall_run['primary_metric_value'])}. These results indicate that once yearly crop choice is coupled to within-season fertilization, the value of hierarchical structure is large enough to dominate the experiment leaderboard. The tradeoff is computational: hierarchical reruns are materially slower than non-hierarchical crop-planning jobs and fertilization runs.

![Runtime tradeoffs](figures/runtime_tradeoffs.png)

*Figure 6. Mean runtime by report group, showing the heavier cost of hierarchical reruns.*

![Runtime comparison](figures/core_runtime_comparison.png)

*Figure 7. Distributional runtime comparison from the core matrix export pack.*

## 4. Ablation results

The 42-run ablation suite should be read as a parameter-isolation study rather than as a second leaderboard. It answers three narrower questions: how much exploration pressure is needed, whether blocked-penalty shaping adds value beyond existing guards, and what nutrient-cost weight behaves most safely across weather regimes.

{markdown_table(list(summary.ablation_takeaways.columns), ablation_rows)}

### 4.1 Entropy coefficient

In fixed weather, the best grouped entropy setting was `ent_coef={summary.point1_fixed['ent_coef']:.2f}` with a mean primary metric of {fmt_num(summary.point1_fixed['primary_metric_value__mean'])}. In random weather, the best grouped setting was `ent_coef={summary.point1_random['ent_coef']:.2f}` with a mean of {fmt_num(summary.point1_random['primary_metric_value__mean'])}. The paired random-weather deterministic-return delta for entropy 0.01 relative to 0.00 was {fmt_num(summary.point1_random_stats['mean_delta'])} with `p={summary.point1_random_stats['p_value']:.4f}`. This is the clearest inferential result in the ablation pack: entropy hurts convergence in deterministic regimes but becomes helpful once weather uncertainty is reintroduced.

![Entropy ablation](figures/ablation_entropy.png)

*Figure 8. Entropy ablation showing the reversal between fixed and random weather.*

### 4.2 Blocked-penalty shaping

The blocked-penalty ablation is descriptive because each configuration is represented by a single seed. Its most defensible finding is not the exact size of the return change, but the mechanism behind it: every recorded season-window compliance value remained 1.0, so the penalty did not improve rule adherence because adherence was already saturated. Instead, it changed the reward landscape. In A2C, moderate or strong penalty values improved the deterministic return. In PPO, the zero-penalty configuration remained best under both weather regimes.

![Blocked penalty](figures/ablation_blocked_penalty.png)

*Figure 9. Deterministic-return comparison across blocked-penalty settings.*

![Compliance](figures/ablation_compliance.png)

*Figure 10. Compliance remained saturated across blocked-penalty settings.*

### 4.3 Nutrient-cost weight

The best fixed-weather cost weight was `{summary.point3_best_fixed['nutrient_cost_weight']:.1f}`, which achieved a mean primary metric of {fmt_num(summary.point3_best_fixed['primary_metric_value__mean'])}. In random weather, the highest grouped mean was observed at weight `{summary.point3_best_random['nutrient_cost_weight']:.1f}`, but the margin over weight 1.0 was only {fmt_num(summary.point3_random_vs_12['mean_delta'])} and was not statistically persuasive (`p={summary.point3_random_vs_12['p_value']:.4f}`). By contrast, under fixed weather the drop from 1.0 to 0.8 was {fmt_num(abs(summary.point3_fixed_vs_08['mean_delta']))}. The academically careful conclusion is therefore that weight 1.0 remains the safest thesis default, while 1.2 may be explored in random weather if a slightly more conservative nutrient policy is desired.

![Cost weight](figures/ablation_cost_weight.png)

*Figure 11. Grouped nutrient-cost-weight performance under fixed and random weather.*

![Cost deltas](figures/ablation_cost_deltas.png)

*Figure 12. Paired delta view for the nutrient-cost-weight ablation.*

## 5. Behavioral diagnostics

Grouped means explain which settings won, but the per-run diagnostics explain how those wins were realized. The extracted diagnostics show that the learning curves in successful runs become progressively smoother, that nutrient application is not uniformly distributed across the season, and that the hierarchical policies produce legible crop-decision traces. These properties matter because they make the experimentation chapter more interpretable than a pure return table.

![Diagnostics panel](figures/diagnostics_panel.png)

*Figure 13. Representative training diagnostics from the blocked-penalty study.*

![Weekly NPK behavior](figures/weekly_npk_behavior.png)

*Figure 14. Weekly NPK behavior from a representative hierarchical run.*

![Crop decision timeline](figures/crop_decision_timeline.png)

*Figure 15. Crop-decision timeline illustrating the policy's yearly planning structure.*

## 6. Discussion and validity boundaries

Three conclusions can be defended cleanly from this evidence pack. First, hierarchical crop planning produces the strongest absolute returns. Second, fertilization behavior is method-sensitive: A2C is highly competitive in deterministic fixed-weather settings, while PPO remains the more robust general choice once broader stochasticity and hierarchy are introduced. Third, the best hyperparameters are regime-dependent: entropy and cost weighting cannot be set intelligently without considering weather uncertainty.

- Simulation is the evaluation medium throughout; none of the conclusions should be presented as field-validated agronomy.
- The active crop-planning and hierarchical experiments remain tied to the current stabilized maize-soy configuration, even though the broader thesis framing is Pakistan-oriented.
- Most grouped comparisons use only three seeds, and the blocked-penalty ablation uses one seed per configuration, so effect size should be interpreted more cautiously than direction of effect.
- DQN was only explored in small reruns and therefore remains a descriptive reference rather than a full competitor in the main ranking.

Within those limits, the practical thesis recommendation is clear: use hierarchical PPO as the strongest overall configuration, keep entropy at 0.00 for fixed weather and 0.01 for random weather, retain nutrient cost weight 1.0 as the default, and apply blocked-penalty shaping only with explicit awareness that its benefit appears algorithm-specific rather than universal.

## 7. Conclusion

The finalized experimentation evidence supports a coherent thesis story. The Pakistan-adapted RL stack is not merely run-ready; it already yields a structured body of results from which method, weather, and design choices can be interpreted with appropriate caution. Hierarchical crop planning is the strongest empirical theme, fixed-weather fertilization remains a competitive niche for A2C, and the ablation suite clarifies which configuration defaults should carry forward into the final thesis narrative. The main scientific value of this chapter lies not only in the scores themselves, but also in the fact that the report can now explain those scores in terms of architecture, stochasticity, reward shaping, and operational behavior.
"""
    (OUTPUT_ROOT / "final_experimentation_report.md").write_text(markdown, encoding="utf-8")


def write_readme() -> None:
    content = textwrap.dedent(
        """
        # Codex Exp Report

        This folder contains a final, thesis-ready experimentation report synthesized from:

        - `Refrence Material`, especially the extracted thesis LaTeX archive used for scope and framing.
        - `final_experiments_report_antigravity`, used as the empirical evidence pack and figure source.
        - The grouped CSV/JSON exports in `artifacts/final_successful_runs/thesis_reporting_pack`, used as the authoritative numeric source.

        ## Main outputs

        - `final_experimentation_report.docx`: the primary deliverable.
        - `final_experimentation_report.md`: a readable source version of the same report.
        - `figures/`: copied evidence figures plus custom summary plots.
        - `diagrams/`: Mermaid sources and rendered diagrams.
        - `tables/`: curated CSV tables extracted for the report.
        - `data/`: copied grouped input tables used during generation.

        ## Rebuild

        Run:

        ```powershell
        python .\\build_report.py
        ```

        Note: the environment used to generate this folder did not include a LaTeX compiler, so the report was emitted in Markdown and Word format rather than as a compiled PDF.
        """
    ).strip()
    (OUTPUT_ROOT / "README.md").write_text(content + "\n", encoding="utf-8")


def main() -> None:
    ensure_dirs()
    copy_reference_figures()
    mermaid_sources = write_mermaid_sources()
    render_mermaid(mermaid_sources)
    save_data_inputs()
    summary = build_summary()
    export_curated_tables(summary)
    create_custom_figures(summary)
    build_docx(summary)
    build_markdown(summary)
    write_readme()
    shutil.copy2(SCRIPT_ROOT / "build_report.py", OUTPUT_ROOT / "build_report.py")
    print(f"Generated report artifacts in: {OUTPUT_ROOT}")


if __name__ == "__main__":
    main()
